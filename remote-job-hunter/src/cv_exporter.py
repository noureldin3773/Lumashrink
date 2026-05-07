from pathlib import Path
import re
import subprocess
from shutil import which
from typing import Optional, Union

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.shared import Inches, Pt
except ImportError:
    Document = None
    WD_ALIGN_PARAGRAPH = None
    WD_LINE_SPACING = None
    WD_TAB_ALIGNMENT = None
    OxmlElement = None
    qn = None
    RELATIONSHIP_TYPE = None
    Inches = None
    Pt = None


CV_DOCX_FILENAME = "Nour Eldin - Senior Product Designer.docx"
CV_PDF_FILENAME = "Nour Eldin - Senior Product Designer.pdf"
DEFAULT_PORTFOLIO_URL = "https://noureldin.framer.ai"
RIGHT_TAB_STOP_INCHES = 7.05


def export_cv_to_docx(
    cv_text: str,
    filename: Union[str, Path] = f"outputs/{CV_DOCX_FILENAME}",
    style: str = "latex",
) -> Path:
    if Document is None or Pt is None or Inches is None:
        raise RuntimeError("python-docx is not installed. Run: pip install -r requirements.txt")

    output_path = Path(filename)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _set_document_style(document, style=style)
    parsed = _parse_cv(cv_text)
    _add_header(document, parsed)
    _add_summary(document, parsed.get("summary", []))
    _add_work_experience(document, parsed.get("work_experience", []))
    _add_skills(document, parsed.get("skills", []))
    _add_education(document, parsed.get("education", []))

    document.save(output_path)
    _export_pdf_if_possible(output_path)
    return output_path


def export_pdf_from_docx(docx_path: Union[str, Path]) -> Optional[Path]:
    path = Path(docx_path)
    if not path.exists():
        return None
    return _export_pdf_if_possible(path)


def _set_document_style(document: "Document", style: str = "latex") -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal_style = document.styles["Normal"]
    if style == "latex":
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(10.5)
        normal_style.paragraph_format.space_after = Pt(2)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(10.5)
        normal_style.paragraph_format.space_after = Pt(4)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _parse_cv(cv_text: str) -> dict[str, list]:
    lines = [line.rstrip() for line in cv_text.splitlines()]
    sections: dict[str, list[str]] = {}
    current = "header"
    sections[current] = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line or line == "---":
            continue
        if line.startswith("# "):
            heading = line[2:].strip().upper()
            if heading in {"SUMMARY", "WORK EXPERIENCE", "SKILLS", "EDUCATION"}:
                current = heading.lower().replace(" ", "_")
                sections.setdefault(current, [])
                continue
        sections.setdefault(current, []).append(line)

    return {
        "header": sections.get("header", []),
        "summary": sections.get("summary", []),
        "work_experience": sections.get("work_experience", []),
        "skills": sections.get("skills", []),
        "education": sections.get("education", []),
    }


def _add_header(document: "Document", parsed: dict[str, list]) -> None:
    header_lines = parsed.get("header", [])
    name = "Nour Eldin Abbas"
    role = "Senior Product Designer"
    email = ""
    phone = ""
    location = ""
    portfolio = DEFAULT_PORTFOLIO_URL

    for line in header_lines:
        clean = line.strip()
        lower = clean.lower()
        if lower.startswith("role:"):
            role = clean.split(":", 1)[1].strip() or role
            continue
        if "@" in clean and not email:
            email = clean
            continue
        if re.search(r"\d{7,}", clean) and not phone:
            phone = clean
            continue
        if ("http" in lower or "www." in lower or ".com" in lower) and "@" not in clean:
            portfolio = _normalize_url(clean)
            continue
        if ("cairo" in lower or "egypt" in lower or "remote" in lower) and not location:
            location = clean
            continue
        if clean and not clean.startswith("#") and clean.lower() != "summary":
            if clean.lower() != role.lower() and "@" not in clean and not re.search(r"\d{7,}", clean):
                name = clean

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(2)
    title_run = title_paragraph.add_run(name)
    title_run.bold = True
    title_run.font.size = Pt(22)

    role_paragraph = document.add_paragraph()
    role_paragraph.paragraph_format.space_after = Pt(4)
    role_run = role_paragraph.add_run(role)
    role_run.italic = True
    role_run.font.size = Pt(13)

    contact_paragraph = document.add_paragraph()
    contact_paragraph.paragraph_format.space_after = Pt(8)
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if email:
        contact_paragraph.add_run(email)
    else:
        contact_paragraph.add_run("Noureldin.ab98@gmail.com")

    plain_parts = [part for part in [phone, location] if part]
    for part in plain_parts:
        contact_paragraph.add_run(" | ")
        contact_paragraph.add_run(part)

    contact_paragraph.add_run(" | ")
    _add_hyperlink(contact_paragraph, "https://www.linkedin.com/in/nour-eldin-ab98-526685145/", "Linkedin")
    contact_paragraph.add_run(" | ")
    _add_hyperlink(contact_paragraph, "https://noureldin.framer.ai/", "Portfolio")


def _add_left_right_line(
    document: "Document",
    left_text: str,
    right_text: str,
    left_size: float = 10.5,
    left_bold: bool = False,
    left_italic: bool = False,
    right_link: str = "",
    after: float = 2,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(RIGHT_TAB_STOP_INCHES), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.paragraph_format.space_after = Pt(after)
    left_run = paragraph.add_run(left_text)
    left_run.font.size = Pt(left_size)
    left_run.bold = left_bold
    left_run.italic = left_italic
    paragraph.add_run("\t")
    if right_link and right_text:
        _add_hyperlink(paragraph, right_link, right_text)
    else:
        right_run = paragraph.add_run(right_text)
        right_run.font.size = Pt(10.5)
        right_run.italic = True


def _normalize_url(value: str) -> str:
    text = value.strip()
    if not text:
        return DEFAULT_PORTFOLIO_URL
    if text.lower().startswith(("http://", "https://")):
        return text
    return f"https://{text.lstrip('/')}"


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    if not url or not text:
        return
    part = paragraph.part
    relationship_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_style = OxmlElement("w:rStyle")
    run_style.set(qn("w:val"), "Hyperlink")
    run_properties.append(run_style)
    run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)


def _add_section_title(document: "Document", title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(9)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"
    divider = document.add_paragraph()
    divider.paragraph_format.space_before = Pt(0)
    divider.paragraph_format.space_after = Pt(3)
    divider_run = divider.add_run("_" * 120)
    divider_run.font.size = Pt(7)


def _add_summary(document: "Document", summary_lines: list[str]) -> None:
    clean_lines = []
    for line in summary_lines:
        text = line.lstrip("- ").strip()
        if text:
            clean_lines.append(text)
    summary_text = " ".join(clean_lines)
    paragraph = document.add_paragraph(summary_text)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def _add_work_experience(document: "Document", lines: list[str]) -> None:
    _add_section_title(document, "WORK EXPERIENCE")
    blocks = _split_experience_blocks(lines)
    sorted_blocks = sorted(blocks, key=lambda block: _experience_start_date(block.get("heading", "")), reverse=True)

    for block in sorted_blocks:
        heading = block.get("heading", "").strip()
        role_text, _ = _split_heading_role_and_dates(heading)
        company_line = block.get("company_line", "").strip()
        intro_text = block.get("intro", "").strip()

        if role_text:
            role_line = document.add_paragraph()
            role_line.paragraph_format.space_before = Pt(4)
            role_line.paragraph_format.space_after = Pt(1)
            role_run = role_line.add_run(role_text)
            role_run.bold = True
            role_run.font.size = Pt(12)

        if company_line:
            company_paragraph = document.add_paragraph()
            company_paragraph.paragraph_format.space_after = Pt(2)
            company_run = company_paragraph.add_run(company_line)
            company_run.font.size = Pt(10.5)
            company_run.italic = True

        if intro_text:
            intro_paragraph = document.add_paragraph(intro_text)
            intro_paragraph.paragraph_format.space_after = Pt(3)
            intro_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        for bullet in block.get("bullets", []):
            item = bullet.lstrip("- ").strip()
            if not item:
                continue
            paragraph = document.add_paragraph(item, style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _split_experience_blocks(lines: list[str]) -> list[dict[str, list[str]]]:
    blocks: list[dict[str, list[str] | str]] = []
    current: dict[str, list[str] | str] = {"heading": "", "company_line": "", "intro": "", "bullets": []}
    consumed_company_line = False
    consumed_intro = False

    for line in lines:
        if line.startswith("## "):
            if current["heading"] or current["bullets"]:
                blocks.append(
                    {
                        "heading": str(current["heading"]),
                        "company_line": str(current.get("company_line", "")),
                        "intro": str(current.get("intro", "")),
                        "bullets": list(current["bullets"]),
                    }
                )
            current = {"heading": line[3:].strip(), "company_line": "", "intro": "", "bullets": []}
            consumed_company_line = False
            consumed_intro = False
        elif line.startswith("- "):
            current["bullets"].append(line)
        elif not consumed_company_line and line.strip():
            current["company_line"] = line.strip()
            consumed_company_line = True
        elif not consumed_intro and line.strip():
            current["intro"] = line.strip()
            consumed_intro = True
        elif line.startswith("- "):
            current["bullets"].append(line)

    if current["heading"] or current["bullets"]:
        blocks.append(
            {
                "heading": str(current["heading"]),
                "company_line": str(current.get("company_line", "")),
                "intro": str(current.get("intro", "")),
                "bullets": list(current["bullets"]),
            }
        )

    return blocks


def _split_heading_role_and_dates(heading: str) -> tuple[str, str]:
    date_match = re.search(r"\(([^)]*)\)", heading)
    if not date_match:
        return heading.strip(), ""
    date_text = date_match.group(1).strip()
    role_text = re.sub(r"\([^)]*\)", "", heading).strip(" -")
    return role_text, date_text


def _experience_start_date(heading: str) -> tuple[int, int]:
    date_match = re.search(r"\(([^)]*)\)", heading)
    if not date_match:
        return (0, 0)
    date_text = date_match.group(1)
    first_date = date_text.split("-", 1)[0].strip()
    month_match = re.search(r"([A-Za-z]{3,9})\s+(\d{4})", first_date)
    year_match = re.search(r"(\d{4})", first_date)
    if not year_match:
        return (0, 0)

    year = int(year_match.group(1))
    month = _month_number(month_match.group(1)) if month_match else 1
    return (year, month)


def _month_number(month: str) -> int:
    months = {
        "jan": 1,
        "feb": 2,
        "mar": 3,
        "apr": 4,
        "may": 5,
        "jun": 6,
        "jul": 7,
        "aug": 8,
        "sep": 9,
        "oct": 10,
        "nov": 11,
        "dec": 12,
    }
    return months.get(month[:3].lower(), 1)


def _add_skills(document: "Document", lines: list[str]) -> None:
    _add_section_title(document, "SKILLS")
    groups: dict[str, list[str]] = {}
    current_heading = "Core"
    groups[current_heading] = []
    for line in lines:
        if line.startswith("## "):
            current_heading = line[3:].strip() or "Skills"
            groups.setdefault(current_heading, [])
            continue
        text = line.lstrip("- ").strip()
        if not text:
            continue
        groups[current_heading].extend([part.strip() for part in text.split(",") if part.strip()])

    normalized_groups = {"Core": [], "Product": [], "Soft": []}
    for heading, items in groups.items():
        h = heading.strip().lower()
        if "core" in h:
            normalized_groups["Core"].extend(items)
        elif "product" in h or "ux" in h or "ui" in h:
            normalized_groups["Product"].extend(items)
        elif "soft" in h or "communication" in h or "collaboration" in h:
            normalized_groups["Soft"].extend(items)
        else:
            normalized_groups["Core"].extend(items)

    priority_order = {
        "figma": 1, "design systems": 2, "product strategy": 3, "ux research": 4, "prototyping": 5,
        "ui design": 6, "ux design": 7, "saas": 8, "ai": 9, "fintech": 10, "b2b": 11,
        "handoff": 12, "usability testing": 13, "communication": 14, "collaboration": 15
    }

    for heading in ["Core", "Product", "Soft"]:
        items = normalized_groups.get(heading, [])
        if not items:
            continue
        heading_paragraph = document.add_paragraph()
        heading_paragraph.paragraph_format.space_before = Pt(4)
        heading_paragraph.paragraph_format.space_after = Pt(2)
        heading_run = heading_paragraph.add_run(heading)
        heading_run.bold = True
        heading_run.font.size = Pt(11)

        unique_items = sorted(
            set(items),
            key=lambda x: (priority_order.get(x.strip().lower(), 999), x.strip().lower()),
        )[:9]
        rows = [unique_items[i : i + 3] for i in range(0, len(unique_items), 3)]
        skills_table = document.add_table(rows=max(1, len(rows)), cols=3)
        skills_table.autofit = True
        for r_index, row_items in enumerate(rows):
            padded = row_items + [""] * (3 - len(row_items))
            for c_index, value in enumerate(padded):
                cell_paragraph = skills_table.cell(r_index, c_index).paragraphs[0]
                cell_paragraph.paragraph_format.space_after = Pt(1.5)
                if value:
                    run = cell_paragraph.add_run(f"• {value}")
                    run.font.size = Pt(10.5)


def _add_education(document: "Document", lines: list[str]) -> None:
    _add_section_title(document, "EDUCATION")
    for line in lines:
        text = line.lstrip("- ").strip()
        if not text:
            continue
        left, _ = _split_education_line(text)
        line_paragraph = document.add_paragraph()
        line_paragraph.paragraph_format.space_after = Pt(1.5)
        line_run = line_paragraph.add_run(left)
        line_run.bold = True


def _split_education_line(text: str) -> tuple[str, str]:
    date_match = re.search(r"\(([^)]*)\)\s*$", text)
    if date_match:
        right = date_match.group(1).strip()
        left = re.sub(r"\([^)]*\)\s*$", "", text).strip(" -")
        return left, right
    dash_match = re.search(r"\s-\s([A-Za-z]{3,9}\s\d{4}.*)$", text)
    if dash_match:
        right = dash_match.group(1).strip()
        left = text[: dash_match.start()].strip()
        return left, right
    return text, ""


def _export_pdf_if_possible(docx_path: Path) -> Optional[Path]:
    pdf_path = docx_path.with_suffix(".pdf")

    try:
        from docx2pdf import convert  # type: ignore

        convert(str(docx_path), str(pdf_path))
        return pdf_path if pdf_path.exists() else None
    except Exception:
        pass

    soffice = which("soffice")
    if not soffice:
        return None

    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
            check=True,
            capture_output=True,
            text=True,
        )
        return pdf_path if pdf_path.exists() else None
    except Exception:
        return None
